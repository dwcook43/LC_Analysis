import glob
import os
import re
import xml.etree.ElementTree as ET

import numpy as np # Numerical library
import matplotlib.pyplot as plt # Plotting library
import pandas as pd # Excel-on-crack library

from docx import Document as doc
from docx.shared import Inches
from adjustText import adjust_text

class LcCsv(object):
    def __init__(self, h5file, ):
            self.store = pd.HDFStore(h5file)

    def folder_proc(self, folder):
        # Just copy and paste the folder name from windows explorer
        # Gotta throw an 'r' in front of this otherwise the '\'s cause problems

        #Note for future revision: add dialog box for user convenience
        # Search string to find wavelengths
        srch2 = re.compile('Sig=(\d+),')
        self.samples=[]
        self.signals=[]
        for file in glob.glob(folder + r'\**\*.D', recursive=True): #changed to .D to capture entire folder
            
            csv_files = glob.glob(file + r'\*.csv', recursive=False)
            if not csv_files:
                print(os.path.basename(file) + ' does not contain any CSVs')
                continue
        
            
            print('Parsing method for file :' + os.path.basename(file))
            #utilize Sample.XML file within .D folder to extract sample name
            tree = ET.parse(file + r'\Sample.XML')
            root = tree.getroot()
            sample = root.find('Name').text
            descrip = root.find('Description') # For Future: Add to metadata in HDF file
            
            infoPath = sample + r'/Method'
            MthdInfo,Detection,Gradient = self.GetMethod(file)
            self.store[infoPath+r'/Wavelengths'] = Detection
            self.store[infoPath + r'/GradientTable'] = Gradient
            self.store[infoPath + r'/MethodInfo'] = MthdInfo
            
            self.samples.append(sample)

            #Now interate through each csv file found in the .D folder
            for signal in csv_files:
                #Determine if it is an integration or signal csv
                if 'Integration' in signal:
                    base = '/integration'
                else:
                    base = '/signal'
                

                # Find the wavelength value from the first line of the file
                with open(signal, encoding='UTF16') as f:
                    firstline = f.readline()
                match2 = srch2.search(firstline)
                
                if match2 is None: #sometimes -- like for MS traces -- there may not be a wavelength
                    wl = firstline.split(',')[0]
                    wl = wl.replace(" ","")
                else:
                    wl = match2.group(1) #Where there is a wavelength, copy it


                path = sample + r'/Signal' + wl + base #Path for HDF file
                
                if (base == '/signal' and wl not in self.signals):
                    self.signals.append(wl)

                
                if path not in self.store: #check to see if already exists
                    #print('Adding:', path)
                    # Put the chromatogram into storage
                    df = pd.read_csv(signal, header=1, encoding='UTF-16 LE')
                    df
                    # Remove start/end whitespace from column names
                    df.rename(columns=lambda x: x.strip(), inplace=True)
                    
                    
                    
                    #add area percent to integration
                    if base == r'/integration':
                        aper = df['Area']/df['Area'].sum()
                        df.insert(3,"Area %",aper,allow_duplicates = False)
                        # may want to format as a percent rather than decimal
                        df.set_index('Peak')

                    self.store[path] = df
                else:
                    print('Skipping:', path)
                

    def select(self, specs, data_type='signal', wl=315):
        if isinstance(specs, str):
            specs = [specs]

        if  data_type== 'signal':
            sep = '/signal'
        elif data_type == 'integration':
            sep = '/integration'
        else:
            raise ValueError("Data type isn't recognized.")

        try:
            wl = str(wl)
        except:
            raise ValueError("Something is wrong with your wavelength"\
                    " selection.")

        suffix = '/Signal' + wl + sep
        dfs = {spec+suffix:self.store[spec+suffix] for spec in specs}
        dfs = pd.concat(dfs, axis=0)
        dfs.index.set_names(['Spec', 'idx'], inplace=True)
        return dfs

    def plotChrom(self,specs,wl=315, default = True, labelRT = True, Integrations = True, figSize = (10,2)):

        specs = [specs]
        
        fig = plt.figure(figsize=figSize)
        fig.set_figheight(2)
        fig.set_figwidth(6)
        dfs = self.select(specs, data_type = 'signal', wl = wl)
        print(dfs.head())
        
        dfs.plot(x = 'Time (min)', y = 'Absorbance (mAu)', legend = False)
        plt.gcf().set_size_inches(figSize)
        if default:
            plt.title(specs)
            plt.xlabel('Time (min)')
            plt.ylabel('Intensity (mAU) @' + str(wl) + ' nm')

            plt.rcParams.update({"axes.facecolor" : "white", 
                             "axes.edgecolor":  "black"})
            plt.autoscale(enable=True, axis='x', tight=True) #autoscale the x axis
         
        if labelRT:
            intData = self.select(specs, data_type = 'integration', wl = wl)
            if not intData():
                print('Integration data does not exist')
                
                texts = []
                x = []
                y = []
                for peaks in intData.index:
                    RT = intData['Retention Time (min)'][peaks]
                    #plt.annotate(f'{RT:.3f}',xy= (RT,x['Height (mAu)'][peaks]),
                    #            xytext=(RT-2,x['Height (mAu)'][peaks]),
                    #            arrowprops=dict(facecolor='black', shrink=0.01))
                    x.append(RT)
                    y.append(intData['Height (mAu)'][peaks])
                    texts.append(plt.text(RT,intData['Height (mAu)'][peaks],f'{RT:.3f}', rotation = 90))
                    if Integrations:
                        bs_slope = (intData['EndIntensity'][peaks]-intData['StartIntensity'][peaks])/(intData['End'][peaks]-intData['Start'][peaks])
                        if bs_slope == 0:
                            b = 0
                        else:
                            b = intData['StartIntensity'][peaks]-bs_slope*intData['Start'][peaks]
                        ind = dfs['Time (min)'].between(intData['Start'][peaks],intData['End'][peaks])
                        x_rt = dfs['Time (min)'][ind]
                        y1 = dfs[ind]['Absorbance (mAu)']
                        y2 = bs_slope*x_rt+b
                       
                        plt.fill_between(x_rt, y1, y2)
            
            adjust_text(texts,x,y,only_move = {'text':'x', 'points': 'x'}, arrowprops=dict(arrowstyle='-'),
                       expand_points = (1.5,3))
                        
    def GetMethod(self,file):
        rx_dict = {
        'Method Name': re.compile(r'Acq\.\sMethod:\s*(.*)\.M'),
        'Flow rate': re.compile(r'Flow:\s*(\d*\.\d*)\smL/min'),
        'Injection volume': re.compile(r'Injection Volume:\s*(\d*.\d*)\s.L'),
        'Column temp': re.compile(r'Temperature:\s*(\d\d\.\d)\s.C'),        
        'Posttime': re.compile(r'Posttime:\s*(\d*.\d*)\smin'),
        'Detector wavelengths': re.compile(r'Yes\sSignal\s([A-Z])\s*(\d\d\d)\s*(\d*)\s*(.*)'),
        'Gradient table': re.compile(r'\s*min\s*%\s*%')
        }
    
        

        def _parse_line(line):
            for key, rx in rx_dict.items():
                match = rx.search(line)
                if  match:
                    return key, match
            return None, None


        def parse_file(filepath):

            filepath = filepath +r'\acq.txt'
            data = []
            Det = pd.DataFrame(columns=['Signal','Wavelength', 'Bandwidth', 'Ref'])
            lct = 1

            with open(filepath, 'r',encoding="utf-16") as file_object:
                line = file_object.readline()

                while line: #read txt file line by line looking for specific parameters
                    key,match = _parse_line(line)

                    if key == 'Method Name':
                        MthdName = match.group(1)
                        
                    if key =='Flow rate':
                        Flowrate = float(match.group(1))
                        
                    if key =='Injection volume':
                        InjVol = float(match.group(1))
                
                    if key =='Posttime':
                        Posttime = float(match.group(1))

                    if key =='Column temp':
                        ColTemp = float(match.group(1))

                    if key =='Detector wavelengths':
                        Signal = match.group(1)
                        DetWave = int(match.group(2))
                        BW = int(match.group(3))
                        refW= (match.group(4))
                        temp = pd.DataFrame([[Signal,DetWave,BW,refW]],columns=['Signal','Wavelength', 'Bandwidth', 'Ref'])
                        Det = Det.append(temp)

                    if key == 'Gradient table':
                        
                        gstart=lct
                        while line.strip(): #counting number of rows
                            line = file_object.readline()
                            lct=lct+1
                        nrows = lct - gstart -2
                        gtable = pd.read_csv(filepath,skiprows=gstart+1, nrows = 3, 
                             header = None, delimiter = '\t', encoding = 'UTF-16')
                        gtable[0] = gtable[0].str.strip()
                        gtable = gtable[0].str.split('\s+', expand=True)
                        if len(gtable.columns) == 7:
                            gtable.columns = ['Time (min)','%A','%B','%C','%D','Flow','Pressure']
                            gtable = gtable.set_index('Time (min)')
                        if len(gtable.columns) == 5:
                            gtable.columns = ['Time (min)','%A','%B','Flow','Pressure']
                            gtable = gtable.set_index('Time (min)')


           
                    line = file_object.readline()
                    lct = lct + 1

                    
                row = {
                    'Method Name': MthdName,
                    'Flow rate (mL/min)': Flowrate,
                    'Method Injection volume (\u03BCL)': InjVol,
                    'Column temp (\u00B0C)': ColTemp,
                    'Posttime (min)': Posttime
                }
                
                data = pd.Series(row,name='Value')
                data = data.to_frame()   
                Det = Det.set_index('Signal')
                

            return data, Det, gtable
        data, Det, gtable = parse_file(file)
        return data, Det, gtable
        
        
    def PlotAll(self, path = 'Default', sample = 'All', figSize = (10,2)):
               
        if sample == 'All':
            samples = self.samples
            #print('All Samples')
        else:
            samples = [sample]
            #print('Selected Sample')
        
        if path == 'Default':
            path = os.getcwd() + '\\Chromatograms'
            
        if not os.path.exists(path):
            os.makedirs(path)
            
        for sample in samples:
            print(sample)
            for signal in self.signals:
                #print(signal)
                self.plotChrom(sample, wl=signal)
                plt.gcf().set_size_inches(figSize)
                plt.savefig(path + '\\' + sample + '_' + signal + r'.png')
                plt.close()
    
    def MakeReport(self, sample = 'All', path = 'Default', docName = 'Default'):

        def doctable(document, data, tabletitle):
            document.add_heading(tabletitle)
            table = document.add_table(rows=(data.shape[0]+1), cols=data.shape[1]+1)  # First row are table headers!
            ind = data.index
            col = data.columns.values.tolist()
            if ind.name is None: ind.name = ' ' #incase there is no index header
            header = [ind.name] + col
            for i, column in enumerate(data) :
                table.cell(0,i).text = header[i] #column headers
                for r, row in enumerate(data.index.values) :
                    table.cell(r+1,0).text = str(ind[r]) #row names
                    if isinstance(data[column][row], np.float64):
                        dvalue = '%.2f' % data[column][row]
                    else:
                        dvalue = str(data[column][row])
                    table.cell(r+1, i+1).text = dvalue
                    
            table.cell(0,data.shape[1]).text = header[-1] #to get last column name
            table.style = 'Light Grid Accent 1'
        
        if sample == 'All':
            samples = self.samples
        else:
            samples = [sample]
        
        for sample in samples:
            if docName == 'Default':
                dName = sample + '_report.docx'
            if path == 'Default':
                path = os.getcwd() + '\\Report'
                print(path)
            self.PlotAll(sample = sample, path = path + '\\Chromatograms')
             
            document = doc()

            document.add_heading('Report',0)

            p = document.add_paragraph('Sample Name: ')
            p.add_run(sample).bold = True

            doctable(document, self.store['/' + sample + '/Method/MethodInfo'], 'Method Information')
            doctable(document, self.store['/' + sample + '/Method/GradientTable'], 'Gradient Table')
            doctable(document, self.store['/' + sample + '/Method/Wavelengths'], 'Signals')

            for signal in self.signals:
                document.add_heading('Signal: ' + signal ,level = 2)
                try: 
                    document.add_picture(path + '\\Chromatograms' + '\\' + sample + r'_' + signal + '.png')
                except KeyError:
                    document.add_paragraph('Signal does not exist in file')
                #print(path + '\\' + samp + r'_' + signal + '.png')
                try:
                    doctable(document, self.store['/' + sample + '/' + 'Signal' + signal + '/integration'], 'Integration Table')
                except:
                    document.add_paragraph('Integration table does not exist in file')
            print(path + '\\' + dName)
            document.save(path + '\\' + dName)


    
            
    def close(self,):
        self.store.close()
        
    def __contains__(self, item):
        return item in self.store